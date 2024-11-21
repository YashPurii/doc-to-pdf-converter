from docx import Document
from reportlab.pdfgen import canvas
from reportlab.lib.pagesizes import letter
from PyPDF2 import PdfReader, PdfWriter
import os

# Metadata Extraction
def extract_metadata(docx_file_path):
    document = Document(docx_file_path)
    core_properties = document.core_properties

    metadata = {
        "title": core_properties.title,
        "author": core_properties.author,
        "created": core_properties.created.strftime('%Y-%m-%d %H:%M:%S') if core_properties.created else None,
        "last_modified": core_properties.modified.strftime('%Y-%m-%d %H:%M:%S') if core_properties.modified else None,
        "subject": core_properties.subject,
        "keywords": core_properties.keywords
    }

    return metadata

# Conversion to PDF
def convert_docx_to_pdf(docx_file_path, output_pdf_path, password=None):
    document = Document(docx_file_path)
    pdf_canvas = canvas.Canvas(output_pdf_path, pagesize=letter)
    width, height = letter

    y_position = height - 50
    for paragraph in document.paragraphs:
        if y_position < 50:
            pdf_canvas.showPage()
            y_position = height - 50
        pdf_canvas.drawString(50, y_position, paragraph.text)
        y_position -= 15

    pdf_canvas.save()

    if password:
        secure_pdf(output_pdf_path, password)

    return output_pdf_path


def secure_pdf(pdf_path, password):
    reader = PdfReader(pdf_path)
    writer = PdfWriter()

    for page in reader.pages:
        writer.add_page(page)

    writer.encrypt(password)

    secure_pdf_path = pdf_path.replace(".pdf", "_secured.pdf")
    with open(secure_pdf_path, "wb") as secured_file:
        writer.write(secured_file)

    os.replace(secure_pdf_path, pdf_path)

# if __name__ == "__main__":
#     docx_path = "C:/Users/yashp/Downloads/Hello.docx"  # Replace with your .docx file path
#     output_pdf = "C:/Users/yashp/Downloads/Hello.pdf"  # Replace with your desired PDF output path
#     password = "mypassword"     # Optional: Set a password

#     try:
#         pdf_path = convert_docx_to_pdf(docx_path, output_pdf, password)
#         print(f"PDF created successfully: {pdf_path}")
#     except Exception as e:
#         print(f"Error: {e}")
